import io
import json
import os
import pathlib
import re
import secrets
import threading
import urllib.parse
import urllib.request

from flask import Flask, jsonify, request, redirect, session, url_for
from flask_cors import CORS
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build

app = Flask(__name__)
app.secret_key = "paiks-dev-secret-key-change-in-prod"
CORS(app, resources={r"/*": {"origins": "*"}}, allow_headers=["Content-Type", "Authorization"], methods=["GET", "POST", "OPTIONS"])

@app.after_request
def _add_cors(response):
    response.headers["Access-Control-Allow-Origin"] = "*"
    response.headers["Access-Control-Allow-Headers"] = "Content-Type, Authorization"
    response.headers["Access-Control-Allow-Methods"] = "GET, POST, OPTIONS"
    return response

@app.errorhandler(Exception)
def _handle_exception(e):
    import traceback
    app.logger.error("Unhandled exception: %s", traceback.format_exc())
    from flask import jsonify
    response = jsonify({"error": str(e)})
    response.status_code = 500
    response.headers["Access-Control-Allow-Origin"] = "*"
    response.headers["Access-Control-Allow-Headers"] = "Content-Type, Authorization"
    response.headers["Access-Control-Allow-Methods"] = "GET, POST, OPTIONS"
    return response

BASE_DIR = pathlib.Path(__file__).resolve().parent
TOKEN_PATH = BASE_DIR / "token.json"
CREDENTIALS_PATH = BASE_DIR / "credentials.json"
SYNC_CACHE_PATH = BASE_DIR / "drive_cache.json"
FOLDER_CONFIG_PATH = BASE_DIR / "folder_config.json"
CHROMA_PATH = BASE_DIR / "chroma_db"
LLM_CONFIG_PATH = BASE_DIR / "llm_config.json"
LOCAL_FILES_PATH = BASE_DIR / "local_files"
LOCAL_FILES_PATH.mkdir(exist_ok=True)

LOCAL_ALLOWED_EXTENSIONS = {".txt", ".md", ".csv", ".docx", ".pdf"}
LOCAL_FILES_CACHE = BASE_DIR / "local_files_cache.json"

# Default: Ollama running locally
_DEFAULT_LLM_CONFIG = {
    "base_url": "http://localhost:11434",
    "model": "llama3.2",
    "provider": "ollama",        # "ollama" | "openai_compat"
}

SCOPES = [
    "https://www.googleapis.com/auth/drive.readonly",
    "https://www.googleapis.com/auth/drive.metadata.readonly",
]

OAUTH_REDIRECT_URI = "http://127.0.0.1:5001/auth/callback"

# Only these MIME types will be synced, browsed, and searched.
ALLOWED_MIME_TYPES = [
    "text/plain",                                                                 # .txt
    "text/csv",                                                                   # .csv
    "application/vnd.google-apps.document",                                      # Google Docs
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",   # .docx
    "application/msword",                                                         # .doc (legacy)
]

# Pre-built Drive API mimeType OR clause, e.g.:
#   (mimeType = 'text/plain' or mimeType = '...' or ...)
_MIME_FILTER = "(" + " or ".join(f"mimeType = '{m}'" for m in ALLOWED_MIME_TYPES) + ")"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _load_client_config():
    """Load client_id and client_secret from credentials.json."""
    data = json.loads(CREDENTIALS_PATH.read_text(encoding="utf-8"))
    # Handle both 'web' and 'installed' credential types
    cfg = data.get("web") or data.get("installed") or {}
    return cfg.get("client_id"), cfg.get("client_secret")


def _get_creds():
    """Return valid Credentials or None."""
    if not TOKEN_PATH.exists():
        return None
    creds = Credentials.from_authorized_user_file(str(TOKEN_PATH), SCOPES)
    if creds and creds.valid:
        return creds
    if creds and creds.expired and creds.refresh_token:
        from google.auth.transport.requests import Request
        creds.refresh(Request())
        TOKEN_PATH.write_text(creds.to_json())
        return creds
    return None


def _drive_service(creds):
    return build("drive", "v3", credentials=creds)


def _fetch_drive_user(creds):
    """Return Google account display name, email, and photo from Drive about API."""
    try:
        service = _drive_service(creds)
        about = service.about().get(fields="user").execute()
        u = about.get("user") or {}
        return {
            "display_name": (u.get("displayName") or "").strip(),
            "email": (u.get("emailAddress") or "").strip(),
            "photo_url": (u.get("photoLink") or "").strip(),
        }
    except Exception:
        return {"display_name": "", "email": "", "photo_url": ""}


def _load_cache():
    if SYNC_CACHE_PATH.exists():
        return json.loads(SYNC_CACHE_PATH.read_text(encoding="utf-8"))
    return {"files": [], "synced_at": None}


def _save_cache(data):
    SYNC_CACHE_PATH.write_text(json.dumps(data, default=str), encoding="utf-8")


def _load_folder_config():
    """Return {"folder_id": ..., "folder_name": ...} or None if not set."""
    if FOLDER_CONFIG_PATH.exists():
        try:
            return json.loads(FOLDER_CONFIG_PATH.read_text(encoding="utf-8"))
        except Exception:
            pass
    return None


def _save_folder_config(folder_id, folder_name):
    FOLDER_CONFIG_PATH.write_text(
        json.dumps({"folder_id": folder_id, "folder_name": folder_name}),
        encoding="utf-8",
    )


def _load_llm_config() -> dict:
    if LLM_CONFIG_PATH.exists():
        try:
            return json.loads(LLM_CONFIG_PATH.read_text(encoding="utf-8"))
        except Exception:
            pass
    return dict(_DEFAULT_LLM_CONFIG)


def _save_llm_config(cfg: dict):
    LLM_CONFIG_PATH.write_text(json.dumps(cfg, indent=2), encoding="utf-8")


def _ollama_list_models(base_url: str) -> list[str]:
    """Return model names available in a running Ollama instance."""
    url = base_url.rstrip("/") + "/api/tags"
    req = urllib.request.Request(url, headers={"Accept": "application/json"})
    with urllib.request.urlopen(req, timeout=5) as resp:
        data = json.loads(resp.read())
    return [m["name"] for m in data.get("models", [])]


def _call_local_llm(prompt: str, cfg: dict) -> str:
    """
    Send a prompt to the configured local LLM and return the text response.
    Supports:
      - Ollama native API  (provider == "ollama" or port 11434)
      - OpenAI-compatible  (provider == "openai_compat", e.g. LM Studio)
    """
    base_url = cfg.get("base_url", "http://localhost:11434").rstrip("/")
    model = cfg.get("model", "llama3.2")
    provider = cfg.get("provider", "ollama")

    payload: dict
    endpoint: str

    if provider == "ollama":
        endpoint = f"{base_url}/api/generate"
        payload = {
            "model": model,
            "prompt": prompt,
            "stream": False,
            "options": {
                "num_predict": 300,
                "temperature": 0.3,
                "num_ctx": 4096,
            },
        }
    else:
        # OpenAI-compatible (LM Studio, llama.cpp server, etc.)
        endpoint = f"{base_url}/v1/chat/completions"
        payload = {
            "model": model,
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.3,
            "max_tokens": 200,
            "stream": False,
        }

    req = urllib.request.Request(
        endpoint,
        data=json.dumps(payload).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=300) as resp:
        result = json.loads(resp.read())

    if provider == "ollama":
        return result.get("response", "").strip()
    else:
        return result["choices"][0]["message"]["content"].strip()


def _build_rag_prompt(query: str, chunks: list[dict]) -> str:
    """Build a compact RAG prompt — shorter prompt = faster inference."""
    ctx_parts = []
    for i, c in enumerate(chunks, 1):
        text = c['text'][:600].strip()
        ctx_parts.append(f"[{i}] {text}")
    context = "\n".join(ctx_parts)
    return (
        f"Based on these documents, answer the question helpfully. Cite sources as [1], [2], etc.\n\n"
        f"{context}\n\n"
        f"Q: {query}\nA:"
    )


# ---------------------------------------------------------------------------
# Health
# ---------------------------------------------------------------------------

@app.get("/health")
def health_check():
    return jsonify({"status": "ok", "service": "flask-api"})


# ---------------------------------------------------------------------------
# Auth (manual OAuth 2.0 — no PKCE)
# ---------------------------------------------------------------------------

@app.get("/auth/status")
def auth_status():
    creds = _get_creds()
    if not creds:
        return jsonify({"authenticated": False, "user": None})
    user = _fetch_drive_user(creds)
    return jsonify({"authenticated": True, "user": user})


@app.get("/auth/url")
def auth_url():
    if not CREDENTIALS_PATH.exists():
        return jsonify({"error": "credentials.json not found. Download it from Google Cloud Console."}), 400
    client_id, _ = _load_client_config()
    if not client_id:
        return jsonify({"error": "Invalid credentials.json format."}), 400

    state = secrets.token_urlsafe(24)
    session["oauth_state"] = state

    params = {
        "client_id": client_id,
        "redirect_uri": OAUTH_REDIRECT_URI,
        "response_type": "code",
        "scope": " ".join(SCOPES),
        "access_type": "offline",
        "prompt": "consent",
        "state": state,
    }
    auth_uri = "https://accounts.google.com/o/oauth2/v2/auth?" + urllib.parse.urlencode(params)
    return jsonify({"url": auth_uri})


@app.get("/auth/callback")
def auth_callback():
    """Receives the auth code forwarded from the Apache bridge."""
    code = request.args.get("code")
    if not code:
        return jsonify({"error": "Missing authorization code"}), 400

    client_id, client_secret = _load_client_config()

    # Exchange auth code for tokens via direct POST (no PKCE needed)
    token_data = urllib.parse.urlencode({
        "code": code,
        "client_id": client_id,
        "client_secret": client_secret,
        "redirect_uri": OAUTH_REDIRECT_URI,
        "grant_type": "authorization_code",
    }).encode("utf-8")

    req = urllib.request.Request(
        "https://oauth2.googleapis.com/token",
        data=token_data,
        headers={"Content-Type": "application/x-www-form-urlencoded"},
    )
    try:
        with urllib.request.urlopen(req) as resp:
            token_response = json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as e:
        error_body = e.read().decode("utf-8")
        app.logger.error("Token exchange failed: %s", error_body)
        return jsonify({"error": "Token exchange failed", "details": json.loads(error_body)}), 400

    # Save token in the format google.oauth2.credentials expects
    token_info = {
        "token": token_response["access_token"],
        "refresh_token": token_response.get("refresh_token"),
        "token_uri": "https://oauth2.googleapis.com/token",
        "client_id": client_id,
        "client_secret": client_secret,
        "scopes": SCOPES,
    }
    TOKEN_PATH.write_text(json.dumps(token_info), encoding="utf-8")

    # Redirect back to Django home
    return redirect("http://127.0.0.1:8000/?connected=1")


@app.post("/auth/disconnect")
def auth_disconnect():
    if TOKEN_PATH.exists():
        TOKEN_PATH.unlink()
    if SYNC_CACHE_PATH.exists():
        SYNC_CACHE_PATH.unlink()
    return jsonify({"status": "disconnected"})


# ---------------------------------------------------------------------------
# Drive
# ---------------------------------------------------------------------------

@app.get("/drive/folders")
def drive_list_folders():
    """Return all folders the user owns/can see so they can pick one."""
    creds = _get_creds()
    if not creds:
        return jsonify({"error": "Not authenticated"}), 401

    try:
        service = _drive_service(creds)
        folders = []
        page_token = None
        while True:
            params = {
                "pageSize": 100,
                "fields": "nextPageToken, files(id, name, modifiedTime)",
                "q": "mimeType = 'application/vnd.google-apps.folder' and trashed = false",
                "orderBy": "name",
            }
            if page_token:
                params["pageToken"] = page_token
            results = service.files().list(**params).execute()
            folders.extend(results.get("files", []))
            page_token = results.get("nextPageToken")
            if not page_token:
                break

        current = _load_folder_config()
        return jsonify({
            "folders": folders,
            "current_folder": current,
        })
    except Exception as e:
        app.logger.error("drive_list_folders error: %s", str(e))
        return jsonify({"error": str(e)}), 500


@app.post("/drive/set-folder")
def drive_set_folder():
    """Save the user's chosen folder. Pass {} to clear (sync all Drive)."""
    payload = request.get_json(silent=True) or {}
    folder_id = payload.get("folder_id", "").strip()
    folder_name = payload.get("folder_name", "").strip()

    if folder_id:
        _save_folder_config(folder_id, folder_name)
        # Wipe stale cache so next sync reflects the new scope
        if SYNC_CACHE_PATH.exists():
            SYNC_CACHE_PATH.unlink()
        return jsonify({"status": "saved", "folder_id": folder_id, "folder_name": folder_name})
    else:
        # Clear folder config → sync entire Drive
        if FOLDER_CONFIG_PATH.exists():
            FOLDER_CONFIG_PATH.unlink()
        if SYNC_CACHE_PATH.exists():
            SYNC_CACHE_PATH.unlink()
        return jsonify({"status": "cleared"})


@app.get("/drive/folder-config")
def drive_folder_config():
    return jsonify(_load_folder_config() or {})


@app.get("/drive/files")
def drive_files():
    creds = _get_creds()
    if not creds:
        return jsonify({"error": "Not authenticated"}), 401

    try:
        page_token = request.args.get("pageToken")
        page_size = int(request.args.get("pageSize", 20))
        name_query = request.args.get("q", "")

        folder_cfg = _load_folder_config()
        service = _drive_service(creds)

        # Build query — scope to chosen folder + allowed file types only
        conditions = ["trashed = false", _MIME_FILTER]
        if folder_cfg and folder_cfg.get("folder_id"):
            conditions.append(f"'{folder_cfg['folder_id']}' in parents")
        if name_query:
            escaped = name_query.replace("'", "\\'")
            conditions.append(f"name contains '{escaped}'")

        params = {
            "pageSize": page_size,
            "fields": "nextPageToken, files(id, name, mimeType, modifiedTime, size, iconLink, webViewLink, thumbnailLink)",
            "orderBy": "modifiedTime desc",
            "q": " and ".join(conditions),
        }
        if page_token:
            params["pageToken"] = page_token

        results = service.files().list(**params).execute()
        return jsonify({
            "files": results.get("files", []),
            "nextPageToken": results.get("nextPageToken"),
            "folder": folder_cfg,
        })
    except Exception as e:
        app.logger.error("drive_files error: %s", str(e))
        return jsonify({"error": str(e)}), 500


@app.post("/drive/sync")
def drive_sync():
    creds = _get_creds()
    if not creds:
        return jsonify({"error": "Not authenticated"}), 401

    folder_cfg = _load_folder_config()
    service = _drive_service(creds)
    all_files = []
    page_token = None

    # Scope the sync to the chosen folder + allowed file types only
    base_q = f"trashed = false and {_MIME_FILTER}"
    if folder_cfg and folder_cfg.get("folder_id"):
        base_q = f"'{folder_cfg['folder_id']}' in parents and trashed = false and {_MIME_FILTER}"

    while True:
        params = {
            "pageSize": 100,
            "fields": "nextPageToken, files(id, name, mimeType, modifiedTime, size, iconLink, webViewLink)",
            "q": base_q,
            "orderBy": "modifiedTime desc",
        }
        if page_token:
            params["pageToken"] = page_token
        results = service.files().list(**params).execute()
        all_files.extend(results.get("files", []))
        page_token = results.get("nextPageToken")
        if not page_token:
            break

    from datetime import datetime, timezone
    cache_data = {
        "files": all_files,
        "synced_at": datetime.now(timezone.utc).isoformat(),
        "total": len(all_files),
        "folder": folder_cfg,
    }
    _save_cache(cache_data)
    return jsonify({
        "status": "synced",
        "total": len(all_files),
        "synced_at": cache_data["synced_at"],
        "folder": folder_cfg,
    })


@app.get("/drive/stats")
def drive_stats():
    creds = _get_creds()
    cache = _load_cache()
    folder_cfg = _load_folder_config()
    file_types = {}
    for f in cache.get("files", []):
        mime = f.get("mimeType", "unknown")
        short = mime.split("/")[-1].split(".")[-1]
        file_types[short] = file_types.get(short, 0) + 1

    return jsonify({
        "authenticated": creds is not None,
        "documents_total": len(cache.get("files", [])),
        "synced_at": cache.get("synced_at", "Not synced yet"),
        "file_types": file_types,
        "folder": folder_cfg,
    })


# ---------------------------------------------------------------------------
# Search
# ---------------------------------------------------------------------------

@app.post("/search")
def search_documents():
    payload = request.get_json(silent=True) or {}
    query = payload.get("query", "").strip()
    if not query:
        return jsonify({"error": "query is required"}), 400

    # Always search cache first — it's instant
    cache = _load_cache()
    q_lower = query.lower()
    words = [w for w in q_lower.split() if w]
    cached_results = [
        {
            "id": f["id"],
            "name": f["name"],
            "mimeType": f.get("mimeType", ""),
            "webViewLink": f.get("webViewLink", ""),
            "modifiedTime": f.get("modifiedTime", ""),
        }
        for f in cache.get("files", [])
        if any(w in f.get("name", "").lower() for w in words)
    ]

    if cached_results:
        return jsonify({"query": query, "results": cached_results[:20], "source": "cache"})

    # Cache miss — try live Drive API search (with a short timeout)
    creds = _get_creds()
    if not creds:
        return jsonify({"query": query, "results": [], "source": "cache"})

    try:
        folder_cfg = _load_folder_config()
        service = _drive_service(creds)
        escaped = query.replace("'", "\\'")
        conditions = [f"name contains '{escaped}'", "trashed = false", _MIME_FILTER]
        if folder_cfg and folder_cfg.get("folder_id"):
            conditions.append(f"'{folder_cfg['folder_id']}' in parents")
        results = (
            service.files()
            .list(
                q=" and ".join(conditions),
                pageSize=20,
                fields="files(id, name, mimeType, modifiedTime, webViewLink, iconLink, size)",
                orderBy="modifiedTime desc",
            )
            .execute()
        )
        return jsonify({"query": query, "results": results.get("files", []), "source": "live"})
    except Exception as e:
        app.logger.warning("Live search failed, returning empty: %s", e)
        return jsonify({"query": query, "results": [], "source": "error"})


# ---------------------------------------------------------------------------
# RAG – ChromaDB with built-in ONNX embeddings (no PyTorch needed)
# ---------------------------------------------------------------------------

_chroma_client = None
_chroma_collection = None
_chroma_lock = threading.Lock()

# Tracks progress of a running ingest so the UI can poll it.
_ingest_progress: dict = {"running": False, "processed": 0, "total": 0, "done": False, "error": None}


def _get_collection():
    global _chroma_client, _chroma_collection
    if _chroma_collection is None:
        with _chroma_lock:
            if _chroma_collection is None:
                import chromadb
                from chromadb.utils.embedding_functions import DefaultEmbeddingFunction
                _chroma_client = chromadb.PersistentClient(path=str(CHROMA_PATH))
                _chroma_collection = _chroma_client.get_or_create_collection(
                    name="paiks_docs",
                    embedding_function=DefaultEmbeddingFunction(),
                    metadata={"hnsw:space": "cosine"},
                )
    return _chroma_collection


def _extract_text_from_drive(service, file_id, mime_type):
    """Download a Drive file and return its plain-text content, or None on failure."""
    try:
        if mime_type == "application/vnd.google-apps.document":
            raw = service.files().export(fileId=file_id, mimeType="text/plain").execute()
            return raw.decode("utf-8", errors="ignore") if isinstance(raw, bytes) else str(raw)

        if mime_type in ("text/plain", "text/csv"):
            from googleapiclient.http import MediaIoBaseDownload
            buf = io.BytesIO()
            dl = MediaIoBaseDownload(buf, service.files().get_media(fileId=file_id))
            done = False
            while not done:
                _, done = dl.next_chunk()
            return buf.getvalue().decode("utf-8", errors="ignore")

        if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            from googleapiclient.http import MediaIoBaseDownload
            import docx as docx_lib
            from lxml import etree
            import zipfile
            buf = io.BytesIO()
            dl = MediaIoBaseDownload(buf, service.files().get_media(fileId=file_id))
            done = False
            while not done:
                _, done = dl.next_chunk()
            buf.seek(0)
            try:
                doc = docx_lib.Document(buf)
                parts = []
                # Extract paragraphs
                for p in doc.paragraphs:
                    if p.text.strip():
                        parts.append(p.text)
                # Extract tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            parts.append(" | ".join(row_text))
                # Extract headers and footers
                for section in doc.sections:
                    for header_p in section.header.paragraphs:
                        if header_p.text.strip():
                            parts.append(header_p.text)
                    for footer_p in section.footer.paragraphs:
                        if footer_p.text.strip():
                            parts.append(footer_p.text)
                # If still empty, extract ALL text from the raw XML (catches text boxes, shapes, etc.)
                if not parts:
                    app.logger.info("python-docx API found no text for %s, trying raw XML extraction", file_id)
                    buf.seek(0)
                    with zipfile.ZipFile(buf) as zf:
                        for name in zf.namelist():
                            if name.endswith('.xml'):
                                tree = etree.parse(zf.open(name))
                                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                                for t_elem in tree.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                                    if t_elem.text and t_elem.text.strip():
                                        parts.append(t_elem.text)
                    if parts:
                        app.logger.info("Raw XML extraction found %d text fragments for %s", len(parts), file_id)
                return "\n".join(parts) if parts else None
            except Exception as exc:
                app.logger.warning("python-docx failed for %s: %s", file_id, exc)
                # Last resort: raw XML extraction from the downloaded bytes
                try:
                    buf.seek(0)
                    import zipfile
                    from lxml import etree
                    parts = []
                    with zipfile.ZipFile(buf) as zf:
                        for name in zf.namelist():
                            if name.endswith('.xml'):
                                tree = etree.parse(zf.open(name))
                                for t_elem in tree.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                                    if t_elem.text and t_elem.text.strip():
                                        parts.append(t_elem.text)
                    if parts:
                        app.logger.info("Raw XML fallback found %d text fragments for %s", len(parts), file_id)
                        return "\n".join(parts)
                except Exception:
                    pass
                app.logger.warning("All extraction methods failed for %s", file_id)
                return None

        if mime_type == "application/msword":
            # Legacy .doc files — try exporting via Google Docs conversion
            try:
                raw = service.files().export(fileId=file_id, mimeType="text/plain").execute()
                return raw.decode("utf-8", errors="ignore") if isinstance(raw, bytes) else str(raw)
            except Exception:
                app.logger.warning("Legacy .doc export failed for %s, skipping", file_id)
                return None

    except Exception as exc:
        app.logger.warning("extract_text failed for %s (%s): %s", file_id, mime_type, exc)
    return None


def _split_sentences(text):
    """Split text into sentences using simple heuristics."""
    # Split on sentence-ending punctuation followed by whitespace
    parts = re.split(r'(?<=[.!?])\s+', text)
    return [s.strip() for s in parts if s.strip()]


def _chunk_text(text, chunk_size=1000, overlap=200):
    """Split text into overlapping sentence-aware chunks.

    Builds chunks by accumulating whole sentences up to *chunk_size* characters.
    When the next sentence would exceed the limit, the current chunk is saved and
    a new one starts with *overlap* characters of trailing context carried over so
    that important concepts at chunk boundaries are not lost.
    """
    text = " ".join(text.split())
    if not text:
        return []

    sentences = _split_sentences(text)
    if not sentences:
        return []

    chunks = []
    current = ""

    for sent in sentences:
        # If a single sentence is longer than chunk_size, split it by chars
        if len(sent) > chunk_size:
            if current:
                chunks.append(current.strip())
                current = ""
            start = 0
            while start < len(sent):
                end = start + chunk_size
                chunks.append(sent[start:end].strip())
                start = end - overlap
            continue

        if len(current) + len(sent) + 1 > chunk_size:
            chunks.append(current.strip())
            # Carry overlap context from the end of the previous chunk
            if len(current) > overlap:
                current = current[-overlap:].lstrip() + " " + sent
            else:
                current = sent
        else:
            current = (current + " " + sent).strip() if current else sent

    if current.strip():
        chunks.append(current.strip())

    return chunks


# ---------------------------------------------------------------------------
# RAG endpoints
# ---------------------------------------------------------------------------

@app.get("/rag/status")
def rag_status():
    try:
        col = _get_collection()
        count = col.count()
        return jsonify({
            "indexed": count > 0,
            "total_chunks": count,
            "ingest_running": _ingest_progress["running"],
            "ingest_progress": _ingest_progress,
        })
    except Exception:
        return jsonify({"indexed": False, "total_chunks": 0, "ingest_running": False})


@app.post("/rag/ingest")
def rag_ingest():
    """Download each synced file, extract text, chunk, embed, and store in ChromaDB."""
    global _ingest_progress

    if _ingest_progress.get("running"):
        return jsonify({"error": "Ingest already running. Check /rag/status for progress."}), 409

    creds = _get_creds()
    if not creds:
        return jsonify({"error": "Not authenticated. Connect Google Drive first."}), 401

    cache = _load_cache()
    files = cache.get("files", [])
    if not files:
        return jsonify({"error": "No synced files found. Run Sync first, then try again."}), 400

    service = _drive_service(creds)
    col = _get_collection()

    _ingest_progress = {"running": True, "processed": 0, "total": len(files), "done": False, "error": None}

    processed, skipped, errors = 0, 0, []
    total_chunks = 0

    for f in files:
        fid = f.get("id", "")
        fname = f.get("name", "untitled")
        mime = f.get("mimeType", "")

        text = _extract_text_from_drive(service, fid, mime)
        if not text or not text.strip():
            app.logger.warning("Skipped '%s' (id=%s, mime=%s) — empty or no text extracted", fname, fid, mime)
            skipped += 1
            errors.append(f"{fname}: empty text (mime={mime})")
            _ingest_progress["processed"] += 1
            continue

        # Skip binary/garbage text — check that most chars are printable
        printable_ratio = sum(1 for c in text[:500] if c.isprintable() or c in '\n\r\t') / max(len(text[:500]), 1)
        if printable_ratio < 0.85:
            app.logger.warning("Skipped '%s' — text appears to be binary (%.0f%% printable)", fname, printable_ratio * 100)
            skipped += 1
            errors.append(f"{fname}: binary content ({printable_ratio*100:.0f}% printable)")
            _ingest_progress["processed"] += 1
            continue

        chunks = _chunk_text(text)
        if not chunks:
            app.logger.warning("Skipped '%s' — text too short to chunk", fname)
            skipped += 1
            errors.append(f"{fname}: too short to chunk")
            _ingest_progress["processed"] += 1
            continue

        ids = [f"{fid}__chunk__{i}" for i in range(len(chunks))]
        metadatas = [
            {
                "file_id": fid,
                "file_name": fname,
                "mime_type": mime,
                "web_view_link": f.get("webViewLink", ""),
                "modified_time": f.get("modifiedTime", ""),
                "chunk_index": i,
            }
            for i in range(len(chunks))
        ]

        try:
            # Pass raw text; ChromaDB's DefaultEmbeddingFunction (ONNX) embeds automatically
            col.upsert(ids=ids, documents=chunks, metadatas=metadatas)
            processed += 1
            total_chunks += len(chunks)
        except Exception as exc:
            errors.append({"file": fname, "error": str(exc)})
            app.logger.error("Chroma upsert failed for %s: %s", fname, exc)

        _ingest_progress["processed"] += 1

    _ingest_progress.update({"running": False, "done": True})

    return jsonify({
        "status": "ingested",
        "files_processed": processed,
        "files_skipped": skipped,
        "total_chunks": total_chunks,
        "errors": errors[:10],
    })


@app.post("/rag/search")
def rag_search():
    """Semantic retrieval → local LLM generation → return answer + source chunks."""
    payload = request.get_json(silent=True) or {}
    query = payload.get("query", "").strip()
    if not query:
        return jsonify({"error": "query is required"}), 400

    folder_cfg = _load_folder_config()
    llm_cfg = _load_llm_config()

    # ── Semantic retrieval ──────────────────────────────────────────────────
    try:
        col = _get_collection()
        if col.count() > 0:
            n_retrieve = min(10, col.count())
            raw = col.query(
                query_texts=[query],   # ChromaDB's ONNX EF embeds the query automatically
                n_results=n_retrieve,
                include=["documents", "metadatas", "distances"],
            )

            # Top-5 chunks — balanced context for local LLM speed.
            MAX_CONTEXT_CHARS = 2000
            llm_chunks: list[dict] = []
            total_chars = 0
            for doc, meta, dist in zip(
                raw["documents"][0][:5],
                raw["metadatas"][0][:5],
                raw["distances"][0][:5],
            ):
                if total_chars + len(doc) > MAX_CONTEXT_CHARS:
                    break
                llm_chunks.append({"name": meta["file_name"], "text": doc})
                total_chars += len(doc)

            # Deduplicated source list for the UI (best chunk per file)
            seen: dict = {}
            for doc, meta, dist in zip(
                raw["documents"][0],
                raw["metadatas"][0],
                raw["distances"][0],
            ):
                fid = meta.get("file_id") or meta.get("file_name", "unknown")
                score = round(1.0 - float(dist), 3)
                if fid not in seen or score > seen[fid]["score"]:
                    seen[fid] = {
                        "id": fid,
                        "name": meta.get("file_name", "Unknown"),
                        "mimeType": meta.get("mime_type", ""),
                        "webViewLink": meta.get("web_view_link", ""),
                        "modifiedTime": meta.get("modified_time", ""),
                        "snippet": doc[:320].strip(),
                        "score": score,
                        "relevance_hint": f"Semantic match · {round(score * 100)}% relevance",
                    }

            hits = sorted(seen.values(), key=lambda x: x["score"], reverse=True)

            # ── Local LLM generation ────────────────────────────────────────
            answer = None
            answer_model = None
            answer_error = None

            if llm_chunks:
                try:
                    prompt = _build_rag_prompt(query, llm_chunks)
                    answer = _call_local_llm(prompt, llm_cfg)
                    answer_model = llm_cfg.get("model")
                except Exception as llm_exc:
                    answer_error = str(llm_exc)
                    app.logger.warning("Local LLM call failed: %s", llm_exc)

            return jsonify({
                "query": query,
                "answer": answer,
                "answer_model": answer_model,
                "answer_error": answer_error,
                "results": hits,
                "total": len(hits),
                "source": "semantic",
                "indexed": True,
                "folder": folder_cfg,
            })
    except Exception as exc:
        import traceback
        app.logger.error("Semantic search error: %s\n%s", exc, traceback.format_exc())
        return jsonify({
            "query": query,
            "answer": None,
            "answer_error": f"Semantic search failed: {exc}",
            "answer_model": None,
            "results": [],
            "total": 0,
            "source": "error",
            "indexed": True,
            "folder": folder_cfg,
        })

    # ── Keyword fallback (index empty or unavailable) ───────────────────────
    creds = _get_creds()
    if creds:
        service = _drive_service(creds)
        escaped = query.replace("'", "\\'")
        conditions = [f"name contains '{escaped}'", "trashed = false", _MIME_FILTER]
        if folder_cfg and folder_cfg.get("folder_id"):
            conditions.append(f"'{folder_cfg['folder_id']}' in parents")
        try:
            res = (
                service.files()
                .list(
                    q=" and ".join(conditions),
                    pageSize=10,
                    fields="files(id, name, mimeType, modifiedTime, webViewLink, size)",
                    orderBy="modifiedTime desc",
                )
                .execute()
            )
            files = res.get("files", [])
        except Exception:
            files = []
    else:
        files = []

    if not files:
        cache = _load_cache()
        words = [w.lower() for w in query.split() if w]
        scored = [
            (sum(1 for w in words if w in f.get("name", "").lower()), f)
            for f in cache.get("files", [])
        ]
        files = [f for score, f in sorted(scored, key=lambda x: x[0], reverse=True) if score > 0][:10]

    words = [w.lower() for w in query.split() if w]
    hits = [
        {
            "id": f.get("id"),
            "name": f.get("name", ""),
            "mimeType": f.get("mimeType", ""),
            "modifiedTime": f.get("modifiedTime", ""),
            "webViewLink": f.get("webViewLink", ""),
            "size": f.get("size"),
            "snippet": None,
            "score": None,
            "relevance_hint": (
                "Matched: " + ", ".join(w for w in words if w in f.get("name", "").lower())
                or "Drive full-text match"
            ),
        }
        for f in files
    ]

    return jsonify({
        "query": query,
        "answer": None,
        "answer_model": None,
        "answer_error": None,
        "results": hits,
        "total": len(hits),
        "source": "keyword",
        "indexed": False,
        "folder": folder_cfg,
    })


# ---------------------------------------------------------------------------
# Local LLM management
# ---------------------------------------------------------------------------

@app.get("/rag/llm/status")
def rag_llm_status():
    """Check if the configured local LLM is reachable and list available models."""
    cfg = _load_llm_config()
    base_url = cfg.get("base_url", "http://localhost:11434").rstrip("/")
    provider = cfg.get("provider", "ollama")

    models: list[str] = []
    reachable = False
    error_msg = None

    try:
        if provider == "ollama":
            models = _ollama_list_models(base_url)
            reachable = True
        else:
            # OpenAI-compat: hit /v1/models
            req = urllib.request.Request(
                f"{base_url}/v1/models",
                headers={"Accept": "application/json"},
            )
            with urllib.request.urlopen(req, timeout=5) as resp:
                data = json.loads(resp.read())
            models = [m["id"] for m in data.get("data", [])]
            reachable = True
    except Exception as exc:
        error_msg = str(exc)

    return jsonify({
        "reachable": reachable,
        "provider": provider,
        "base_url": base_url,
        "current_model": cfg.get("model"),
        "available_models": models,
        "error": error_msg,
    })


@app.post("/rag/llm/config")
def rag_llm_config():
    """Save the local LLM configuration."""
    payload = request.get_json(silent=True) or {}
    base_url = payload.get("base_url", "").strip()
    model = payload.get("model", "").strip()
    provider = payload.get("provider", "ollama").strip()

    if not base_url or not model:
        return jsonify({"error": "base_url and model are required"}), 400

    cfg = {"base_url": base_url, "model": model, "provider": provider}
    _save_llm_config(cfg)
    return jsonify({"status": "saved", **cfg})


# ---------------------------------------------------------------------------
# Local File helpers
# ---------------------------------------------------------------------------

def _local_files_meta():
    """Return list of local file metadata dicts from cache."""
    try:
        if LOCAL_FILES_CACHE.exists():
            return json.loads(LOCAL_FILES_CACHE.read_text(encoding="utf-8"))
    except Exception:
        pass
    return []

def _local_files_meta_save(meta_list):
    LOCAL_FILES_CACHE.write_text(json.dumps(meta_list, ensure_ascii=False, indent=2), encoding="utf-8")

def _extract_text_from_local(filepath: pathlib.Path) -> str | None:
    """Extract plain text from a local file by extension."""
    ext = filepath.suffix.lower()
    try:
        if ext in (".txt", ".md", ".csv"):
            return filepath.read_text(encoding="utf-8", errors="ignore")

        if ext == ".docx":
            import docx as docx_lib
            doc = docx_lib.Document(str(filepath))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = [c.text.strip() for c in row.cells if c.text.strip()]
                    if row_text:
                        parts.append(" | ".join(row_text))
            return "\n".join(parts) if parts else None

        if ext == ".pdf":
            try:
                import pdfplumber
                with pdfplumber.open(str(filepath)) as pdf:
                    pages = [p.extract_text() or "" for p in pdf.pages]
                return "\n".join(pages).strip() or None
            except ImportError:
                pass
            # Fallback: read raw bytes and decode printable text
            raw = filepath.read_bytes()
            text = raw.decode("latin-1", errors="ignore")
            printable = "".join(c if c.isprintable() or c in "\n\r\t" else " " for c in text)
            return printable.strip() or None

    except Exception as exc:
        app.logger.warning("local extract failed for %s: %s", filepath.name, exc)
    return None


@app.get("/local/files")
def local_files_list():
    """Return list of uploaded local files with metadata."""
    return jsonify({"files": _local_files_meta()})


@app.post("/local/upload")
def local_upload():
    """Accept one or more file uploads, extract text, chunk, and embed into ChromaDB."""
    uploaded = request.files.getlist("files")
    if not uploaded:
        return jsonify({"error": "No files provided"}), 400

    col = _get_collection()
    meta_list = _local_files_meta()
    results = []

    for f in uploaded:
        original_name = f.filename or "upload"
        ext = pathlib.Path(original_name).suffix.lower()
        if ext not in LOCAL_ALLOWED_EXTENSIONS:
            results.append({"name": original_name, "status": "skipped", "reason": f"Unsupported type {ext}"})
            continue

        # Safe filename: strip path components
        safe_name = pathlib.Path(original_name).name
        dest = LOCAL_FILES_PATH / safe_name
        # If duplicate name, append counter
        counter = 1
        while dest.exists():
            stem = pathlib.Path(safe_name).stem
            dest = LOCAL_FILES_PATH / f"{stem}_{counter}{ext}"
            counter += 1

        f.save(str(dest))
        size = dest.stat().st_size

        text = _extract_text_from_local(dest)
        if not text or not text.strip():
            dest.unlink(missing_ok=True)
            results.append({"name": original_name, "status": "error", "reason": "Could not extract text"})
            continue

        chunks = _chunk_text(text)
        if not chunks:
            dest.unlink(missing_ok=True)
            results.append({"name": original_name, "status": "error", "reason": "Text too short to chunk"})
            continue

        file_id = "local__" + dest.name
        # Remove old chunks for this file if re-uploading
        try:
            existing = col.get(where={"source": "local", "file_name": dest.name})
            if existing and existing.get("ids"):
                col.delete(ids=existing["ids"])
        except Exception:
            pass

        ids = [f"{file_id}__chunk__{i}" for i in range(len(chunks))]
        metadatas = [
            {"source": "local", "file_name": dest.name, "file_path": str(dest), "chunk_index": i}
            for i in range(len(chunks))
        ]
        col.upsert(ids=ids, documents=chunks, metadatas=metadatas)

        import time
        file_meta = {
            "id": file_id,
            "name": dest.name,
            "original_name": original_name,
            "size": size,
            "chunks": len(chunks),
            "uploaded_at": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
            "ext": ext,
        }
        # Update or append in meta list
        meta_list = [m for m in meta_list if m.get("id") != file_id]
        meta_list.append(file_meta)
        results.append({"name": original_name, "status": "indexed", "chunks": len(chunks)})

    _local_files_meta_save(meta_list)
    return jsonify({"results": results, "total_files": len([r for r in results if r["status"] == "indexed"])})


@app.post("/local/delete")
def local_delete():
    """Delete a local file and remove its ChromaDB chunks."""
    body = request.get_json(silent=True) or {}
    file_id = body.get("file_id", "")
    if not file_id or not file_id.startswith("local__"):
        return jsonify({"error": "Invalid file_id"}), 400

    file_name = file_id.replace("local__", "", 1)
    dest = LOCAL_FILES_PATH / file_name

    # Remove from ChromaDB
    try:
        col = _get_collection()
        existing = col.get(where={"source": "local", "file_name": file_name})
        if existing and existing.get("ids"):
            col.delete(ids=existing["ids"])
    except Exception as exc:
        app.logger.warning("ChromaDB delete error for %s: %s", file_name, exc)

    # Remove file from disk
    if dest.exists():
        dest.unlink()

    # Remove from meta cache
    meta_list = [m for m in _local_files_meta() if m.get("id") != file_id]
    _local_files_meta_save(meta_list)

    return jsonify({"status": "deleted", "file_id": file_id})


def _prewarm_chroma():
    """Load ChromaDB + ONNX embedding model at startup so first query is fast."""
    try:
        col = _get_collection()
        if col.count() > 0:
            # Warm up the embedding model with a tiny query
            col.query(query_texts=["warmup"], n_results=1, include=["documents"])
        app.logger.info("ChromaDB pre-warmed (%d chunks)", col.count())
    except Exception as e:
        app.logger.warning("ChromaDB pre-warm failed (non-fatal): %s", e)


if __name__ == "__main__":
    os.environ["OAUTHLIB_INSECURE_TRANSPORT"] = "1"  # Allow HTTP for local dev
    _prewarm_chroma()
    app.run(host="127.0.0.1", port=5001, debug=True)
